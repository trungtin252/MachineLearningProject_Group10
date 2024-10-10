# import numpy as np
# import pandas as pd
# import matplotlib.pyplot as plt
# import seaborn as sns
# from sklearn.model_selection import train_test_split
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, confusion_matrix
# from docx import Document
# from docx.shared import Inches
# import io

# # Đọc dữ liệu từ file CSV
# file_path = 'D:\\THBuoi2_NguyenHaiDuong_B2013465\\iris_data.csv'  # Thay đổi đường dẫn nếu cần
# data = pd.read_csv(file_path)

# # Giả sử rằng cột cuối cùng là nhãn lớp
# X = data.iloc[:, :-1].values  # Tất cả các cột trừ cột cuối
# y = data.iloc[:, -1].values   # Cột cuối cùng

# # Hàm để vẽ ma trận phân lớp và lưu vào đối tượng BytesIO
# def save_confusion_matrix_to_image(cm, title='Ma trận phân lớp'):
#     buf = io.BytesIO()
#     plt.figure(figsize=(8, 6))
#     sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=data.columns[:-1], yticklabels=data.columns[:-1])
#     plt.title(title)
#     plt.xlabel('Nhãn dự đoán')
#     plt.ylabel('Nhãn thật')
#     plt.savefig(buf, format='png')
#     buf.seek(0)
#     plt.close()
#     return buf

# # Hàm đánh giá mô hình Naive Bayes
# def naive_bayes_evaluation():
#     # Phân chia dữ liệu theo 80% Train, 20% Test
#     X_train_80, X_test_20, y_train_80, y_test_20 = train_test_split(X, y, test_size=0.2, random_state=42)
    
#     # Phân chia dữ liệu theo phương pháp Hold-out (2/3 Train, 1/3 Test)
#     X_train_66, X_test_33, y_train_66, y_test_33 = train_test_split(X, y, test_size=1/3, random_state=42)

#     # Mô hình Naive Bayes
#     nb = GaussianNB()

#     # Tạo đối tượng Document
#     doc = Document()
#     doc.add_heading('Đánh giá mô hình Naive Bayes', 0)

#     # Đánh giá với 80% Train, 20% Test
#     nb.fit(X_train_80, y_train_80)
#     y_pred_20 = nb.predict(X_test_20)
#     acc_20 = accuracy_score(y_test_20, y_pred_20)
#     doc.add_paragraph(f"Naive Bayes (80% Train, 20% Test): Độ chính xác = {acc_20:.2f}")

#     cm_20 = confusion_matrix(y_test_20, y_pred_20, labels=np.unique(y))
#     buf_20 = save_confusion_matrix_to_image(cm_20, title='Ma trận phân lớp Naive Bayes (80% Train, 20% Test)')
#     doc.add_picture(buf_20, width=Inches(6.0))
    
#     # Đánh giá với Hold-out (2/3 Train, 1/3 Test)
#     nb.fit(X_train_66, y_train_66)
#     y_pred_33 = nb.predict(X_test_33)
#     acc_33 = accuracy_score(y_test_33, y_pred_33)
#     doc.add_paragraph(f"Naive Bayes (Hold-out 2/3 Train, 1/3 Test): Độ chính xác = {acc_33:.2f}")

#     cm_33 = confusion_matrix(y_test_33, y_pred_33, labels=np.unique(y))
#     buf_33 = save_confusion_matrix_to_image(cm_33, title='Ma trận phân lớp Naive Bayes (Hold-out 2/3 Train, 1/3 Test)')
#     doc.add_picture(buf_33, width=Inches(6.0))

#     # Lưu tài liệu
#     file_name = "D:\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.docx"
#     doc.save(file_name)
#     print(f"Kết quả đã được lưu vào file: {file_name}")
# # Thực hiện đánh giá
# naive_bayes_evaluation()



# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import LabelEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import train_test_split
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Lấy thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# imputer.fit(X[:, 3:7])
# X[:, 3:7] = imputer.transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng LabelEncoder
# encoder = LabelEncoder()
# X[:, 3] = encoder.fit_transform(X[:, 3])  # Mã hóa cột 'Vict Sex'
# X[:, 4] = encoder.fit_transform(X[:, 4])  # Mã hóa cột 'Vict Descent'
# X[:, 7] = encoder.fit_transform(X[:, 7])  # Mã hóa cột 'LOCATION'

# # Phân chia dữ liệu
# for split_ratio in [2/3, 0.8]:
#     test_size = 1 - split_ratio
#     X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=42)
    
#     print(f"Số lượng mẫu trong tập Train: {len(X_train)}")
#     print(f"Số lượng mẫu trong tập Test: {len(X_test)}")

#     # Xây dựng mô hình Gaussian Naive Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred = gnb.predict(X_test)

#     # Đánh giá độ chính xác
#     accuracy = accuracy_score(y_test, y_pred) * 100  # Nhân với 100 để chuyển thành %
#     print(f'Phân chia dữ liệu: {int(split_ratio * 100)}% Train, {int(test_size * 100)}% Test')
#     print(f'Độ chính xác tổng thể: {accuracy:.2f}%')

#     # In classification report với zero_division=1 để tránh cảnh báo
#     print('Báo cáo phân loại:')
#     print(classification_report(y_test, y_pred, zero_division=1))


# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import LabelEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import train_test_split, KFold
# from sklearn.neighbors import KNeighborsClassifier
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
# import seaborn as sns
# import matplotlib.pyplot as plt
# import io
# from docx import Document
# from docx.shared import Inches
# import os

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Chọn thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# X[:, 3:7] = imputer.fit_transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng LabelEncoder
# encoder = LabelEncoder()
# X[:, 3] = encoder.fit_transform(X[:, 3])  # Encode 'Vict Sex'
# X[:, 4] = encoder.fit_transform(X[:, 4])  # Encode 'Vict Descent'
# X[:, 7] = encoder.fit_transform(X[:, 7])  # Encode 'LOCATION'

# # Đường dẫn file DOCX
# docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes_KNN.docx'
# if os.path.exists(docx_file_path):
#     os.remove(docx_file_path)
# doc = Document()

# # Hàm để lưu ma trận phân lớp vào file DOCX
# def save_confusion_matrix_to_image(cm, title='Ma trận phân lớp'):
#     buf = io.BytesIO()
#     plt.figure(figsize=(8, 6))
#     sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=np.unique(y), yticklabels=np.unique(y))
#     plt.title(title)
#     plt.xlabel('Nhãn dự đoán')
#     plt.ylabel('Nhãn thực tế')
#     plt.savefig(buf, format='png')
#     buf.seek(0)
#     plt.close()
#     return buf

# # Phân chia dữ liệu
# kf = KFold(n_splits=5, shuffle=True, random_state=42)

# for train_index, test_index in kf.split(X):
#     X_train, X_test = X[train_index], X[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     # Số lượng phần tử và nhãn trong tập test
#     print(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     print(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     doc.add_paragraph(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     doc.add_paragraph(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     # Bước d: Mô hình KNN
#     knn = KNeighborsClassifier(n_neighbors=9)
#     knn.fit(X_train, y_train)
#     y_pred_knn = knn.predict(X_test)

#     # Đánh giá KNN
#     accuracy_knn = accuracy_score(y_test, y_pred_knn) * 100
#     print(f"Độ chính xác tổng thể của KNN: {accuracy_knn:.2f}%")
#     print("Báo cáo phân loại KNN:")
#     print(classification_report(y_test, y_pred_knn, zero_division=1))

#     doc.add_paragraph(f"Độ chính xác tổng thể của KNN: {accuracy_knn:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại KNN:")
#     doc.add_paragraph(classification_report(y_test, y_pred_knn, zero_division=1))

#     # In ra giá trị của 7 phần tử đầu tiên
#     print(f"Giá trị nhãn thực tế của 7 phần tử đầu tiên: {y_test[:7]}")
#     print(f"Giá trị nhãn dự đoán của 7 phần tử đầu tiên: {y_pred_knn[:7]}")
    
#     doc.add_paragraph(f"Giá trị nhãn thực tế của 7 phần tử đầu tiên: {y_test[:7]}")
#     doc.add_paragraph(f"Giá trị nhãn dự đoán của 7 phần tử đầu tiên: {y_pred_knn[:7]}")

#     # Bước d: Ma trận phân lớp KNN
#     cm_knn = confusion_matrix(y_test, y_pred_knn)
#     print("Ma trận phân lớp KNN:")
#     print(cm_knn)

#     buf_cm_knn = save_confusion_matrix_to_image(cm_knn, title='Ma trận phân lớp KNN')
#     doc.add_paragraph("Ma trận phân lớp của KNN:")
#     doc.add_picture(buf_cm_knn, width=Inches(6))

#     # Bước e: Mô hình Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred_bayes = gnb.predict(X_test)

#     # Đánh giá Naive Bayes
#     accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
#     print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     print("Báo cáo phân loại Naive Bayes:")
#     print(classification_report(y_test, y_pred_bayes, zero_division=1))

#     doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
#     doc.add_paragraph(classification_report(y_test, y_pred_bayes, zero_division=1))

#     # Bước e: Ma trận phân lớp Naive Bayes
#     cm_bayes = confusion_matrix(y_test, y_pred_bayes)
#     print("Ma trận phân lớp Naive Bayes:")
#     print(cm_bayes)

#     buf_cm_bayes = save_confusion_matrix_to_image(cm_bayes, title='Ma trận phân lớp Naive Bayes')
#     doc.add_paragraph("Ma trận phân lớp của Naive Bayes:")
#     doc.add_picture(buf_cm_bayes, width=Inches(6))

#     break  # Chỉ chạy lần phân chia đầu tiên

# # Lưu file DOCX
# doc.save(docx_file_path)
# print(f"Kết quả đã được lưu vào file: {docx_file_path}")




# # 26%

# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import LabelEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import KFold
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
# import seaborn as sns
# import matplotlib.pyplot as plt
# import io
# from docx import Document
# from docx.shared import Inches
# import os

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Chọn thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# X[:, 3:7] = imputer.fit_transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng LabelEncoder
# encoder = LabelEncoder()
# X[:, 3] = encoder.fit_transform(X[:, 3])  # Encode 'Vict Sex'
# X[:, 4] = encoder.fit_transform(X[:, 4])  # Encode 'Vict Descent'
# X[:, 7] = encoder.fit_transform(X[:, 7])  # Encode 'LOCATION'

# # Đường dẫn file DOCX
# docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.docx'
# if os.path.exists(docx_file_path):
#     os.remove(docx_file_path)
# doc = Document()

# # Hàm để lưu ma trận phân lớp vào file DOCX
# def save_confusion_matrix_to_image(cm, title='Ma trận phân lớp'):
#     buf = io.BytesIO()
#     plt.figure(figsize=(8, 6))
#     sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=np.unique(y), yticklabels=np.unique(y))
#     plt.title(title)
#     plt.xlabel('Nhãn dự đoán')
#     plt.ylabel('Nhãn thực tế')
#     plt.savefig(buf, format='png')
#     buf.seek(0)
#     plt.close()
#     return buf

# # Phân chia dữ liệu
# kf = KFold(n_splits=5, shuffle=True, random_state=42)

# for train_index, test_index in kf.split(X):
#     X_train, X_test = X[train_index], X[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     # Số lượng phần tử và nhãn trong tập test 
#     print(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     print(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     doc.add_paragraph(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     doc.add_paragraph(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     # Mô hình Naive Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred_bayes = gnb.predict(X_test)

#     # Đánh giá Naive Bayes
#     accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
#     print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     print("Báo cáo phân loại Naive Bayes:")
#     print(classification_report(y_test, y_pred_bayes, zero_division=1))

#     doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
#     doc.add_paragraph(classification_report(y_test, y_pred_bayes, zero_division=1))

#     # Ma trận phân lớp Naive Bayes
#     cm_bayes = confusion_matrix(y_test, y_pred_bayes)
#     print("Ma trận phân lớp Naive Bayes:")
#     print(cm_bayes)

#     buf_cm_bayes = save_confusion_matrix_to_image(cm_bayes, title='Ma trận phân lớp Naive Bayes')
#     doc.add_paragraph("Ma trận phân lớp của Naive Bayes:")
#     doc.add_picture(buf_cm_bayes, width=Inches(6))

#     break  # Chỉ chạy lần phân chia đầu tiên

# # Lưu file DOCX
# doc.save(docx_file_path)
# print(f"Kết quả đã được lưu vào file: {docx_file_path}")





#  33%
 
import pandas as pd
import numpy as np
from sklearn.preprocessing import LabelEncoder
from sklearn.impute import SimpleImputer
from sklearn.model_selection import KFold
from sklearn.naive_bayes import GaussianNB
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
import seaborn as sns
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches
import os

# Đọc dữ liệu
df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# Chọn thuộc tính và nhãn
data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# Chuyển dữ liệu thành dạng array
X = data.iloc[:, :-1].values
y = data.iloc[:, -1].values

# Xử lý những thuộc tính chứa giá trị null
# Tách riêng các cột số và cột phân loại
numeric_columns = [0, 1, 2, 5, 6]  # Các cột số: TIME OCC, AREA, Vict Age, Premis Cd, Weapon Used Cd
categorical_columns = [3, 4, 7]    # Các cột phân loại: Vict Sex, Vict Descent, LOCATION

# Xử lý cột số (chiến lược 'mean' cho các cột số)
numeric_imputer = SimpleImputer(missing_values=np.nan, strategy="mean")
X[:, numeric_columns] = numeric_imputer.fit_transform(X[:, numeric_columns])

# Xử lý cột phân loại (chiến lược 'most_frequent' cho các cột phân loại)
categorical_imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
X[:, categorical_columns] = categorical_imputer.fit_transform(X[:, categorical_columns])

# Mã hóa các thuộc tính phân loại bằng LabelEncoder
encoder = LabelEncoder()
for col in categorical_columns:
    X[:, col] = encoder.fit_transform(X[:, col])

# Đường dẫn file DOCX
docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes2.docx'
if os.path.exists(docx_file_path):
    os.remove(docx_file_path)
doc = Document()

# Hàm để lưu ma trận phân lớp vào file DOCX
def save_confusion_matrix_to_image(cm, title='Ma trận phân lớp'):
    buf = io.BytesIO()
    plt.figure(figsize=(8, 6))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=np.unique(y), yticklabels=np.unique(y))
    plt.title(title)
    plt.xlabel('Nhãn dự đoán')
    plt.ylabel('Nhãn thực tế')
    plt.savefig(buf, format='png')
    buf.seek(0)
    plt.close()
    return buf

# Phân chia dữ liệu
kf = KFold(n_splits=5, shuffle=True, random_state=42)

for train_index, test_index in kf.split(X):
    X_train, X_test = X[train_index], X[test_index]
    y_train, y_test = y[train_index], y[test_index]

    # Số lượng phần tử và nhãn trong tập test
    print(f"Số lượng phần tử trong tập test: {len(X_test)}")
    print(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

    doc.add_paragraph(f"Số lượng phần tử trong tập test: {len(X_test)}")
    doc.add_paragraph(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

    # Mô hình Naive Bayes
    gnb = GaussianNB()
    gnb.fit(X_train, y_train)
    y_pred_bayes = gnb.predict(X_test)

    # Đánh giá Naive Bayes
    accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
    print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
    print("Báo cáo phân loại Naive Bayes:")
    print(classification_report(y_test, y_pred_bayes, zero_division=1))

    doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
    doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
    doc.add_paragraph(classification_report(y_test, y_pred_bayes, zero_division=1))

    # Ma trận phân lớp Naive Bayes
    cm_bayes = confusion_matrix(y_test, y_pred_bayes)
    print("Ma trận phân lớp Naive Bayes:")
    print(cm_bayes)

    buf_cm_bayes = save_confusion_matrix_to_image(cm_bayes, title='Ma trận phân lớp Naive Bayes')
    doc.add_paragraph("Ma trận phân lớp của Naive Bayes:")
    doc.add_picture(buf_cm_bayes, width=Inches(6))

    break  # Chỉ chạy lần phân chia đầu tiên

# Lưu file DOCX
doc.save(docx_file_path)
print(f"Kết quả đã được lưu vào file: {docx_file_path}")




# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import LabelEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import KFold
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report
# import os
# from docx import Document

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Chọn thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# X[:, 3:7] = imputer.fit_transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng LabelEncoder
# encoder = LabelEncoder()
# X[:, 3] = encoder.fit_transform(X[:, 3])  # Encode 'Vict Sex'
# X[:, 4] = encoder.fit_transform(X[:, 4])  # Encode 'Vict Descent'
# X[:, 7] = encoder.fit_transform(X[:, 7])  # Encode 'LOCATION'

# # Đường dẫn file DOCX
# docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.docx'
# if os.path.exists(docx_file_path):
#     os.remove(docx_file_path)
# doc = Document()

# # Phân chia dữ liệu
# kf = KFold(n_splits=5, shuffle=True, random_state=42)

# for train_index, test_index in kf.split(X):
#     X_train, X_test = X[train_index], X[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     # Số lượng phần tử và nhãn trong tập test 
#     print(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     print(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     doc.add_paragraph(f"Số lượng phần tử trong tập test: {len(X_test)}")
#     doc.add_paragraph(f"Giá trị nhãn khác nhau trong tập test: {np.unique(y_test)}")

#     # Mô hình Naive Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred_bayes = gnb.predict(X_test)

#     # Đánh giá Naive Bayes
#     accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
#     print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     print("Báo cáo phân loại Naive Bayes:")
#     print(classification_report(y_test, y_pred_bayes, zero_division=1))

#     doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
#     doc.add_paragraph(classification_report(y_test, y_pred_bayes, zero_division=1))

#     break  # Chỉ chạy lần phân chia đầu tiên

# # Lưu file DOCX
# doc.save(docx_file_path)
# print(f"Kết quả đã được lưu vào file: {docx_file_path}")




# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import OneHotEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import KFold
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report
# from docx import Document
# import os

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Chọn thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# X[:, 3:7] = imputer.fit_transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng OneHotEncoder
# encoder = OneHotEncoder(drop='first', sparse=False)
# categorical_columns = [3, 4, 7]  # Index của các cột cần mã hóa
# X_encoded = encoder.fit_transform(X[:, categorical_columns])

# # Kết hợp với các thuộc tính số khác
# X = np.concatenate([X[:, :3], X_encoded], axis=1)

# # Đường dẫn file DOCX
# docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.docx'
# if os.path.exists(docx_file_path):
#     os.remove(docx_file_path)
# doc = Document()

# # Phân chia dữ liệu
# kf = KFold(n_splits=5, shuffle=True, random_state=42)

# for train_index, test_index in kf.split(X):
#     X_train, X_test = X[train_index], X[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     # Mô hình Naive Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred_bayes = gnb.predict(X_test)

#     # Đánh giá Naive Bayes
#     accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
#     print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     print("Báo cáo phân loại Naive Bayes:")
#     print(classification_report(y_test, y_pred_bayes, zero_division=1))

#     doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
#     doc.add_paragraph(classification_report(y_test, y_pred_bayes, zero_division=1))

#     break  # Chỉ chạy lần phân chia đầu tiên

# # Lưu file DOCX
# doc.save(docx_file_path)
# print(f"Kết quả đã được lưu vào file: {docx_file_path}")


# import pandas as pd
# import numpy as np
# from sklearn.preprocessing import LabelEncoder
# from sklearn.impute import SimpleImputer
# from sklearn.model_selection import KFold
# from sklearn.naive_bayes import GaussianNB
# from sklearn.metrics import accuracy_score, classification_report
# from docx import Document
# import os

# # Đọc dữ liệu
# df = pd.read_csv("D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\Crime_Data_from_2020_to_Present.csv", sep=",")

# # Chọn thuộc tính và nhãn
# data = df[['TIME OCC', 'AREA', 'Vict Age', 'Vict Sex', 'Vict Descent', 'Premis Cd', 'Weapon Used Cd', 'LOCATION', 'Crm Cd']]

# # Chuyển dữ liệu thành dạng array
# X = data.iloc[:, :-1].values
# y = data.iloc[:, -1].values

# # Xử lý những thuộc tính chứa giá trị null
# imputer = SimpleImputer(missing_values=np.nan, strategy="most_frequent")
# X[:, 3:7] = imputer.fit_transform(X[:, 3:7])

# # Mã hóa dữ liệu bằng LabelEncoder
# encoder = LabelEncoder()
# X[:, 3] = encoder.fit_transform(X[:, 3])  # Encode 'Vict Sex'
# X[:, 4] = encoder.fit_transform(X[:, 4])  # Encode 'Vict Descent'
# X[:, 7] = encoder.fit_transform(X[:, 7])  # Encode 'LOCATION'

# # Đường dẫn file DOCX
# docx_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.docx'
# if os.path.exists(docx_file_path):
#     os.remove(docx_file_path)
# doc = Document()

# # Đường dẫn file TXT
# txt_file_path = 'D:\\MHUD\\THBuoi2_NguyenHaiDuong_B2013465\\ThongKeKQ_Bayes.txt'
# if os.path.exists(txt_file_path):
#     os.remove(txt_file_path)

# # Phân chia dữ liệu
# kf = KFold(n_splits=5, shuffle=True, random_state=42)

# for train_index, test_index in kf.split(X):
#     X_train, X_test = X[train_index], X[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     # Số lượng phần tử và nhãn trong tập test 
#     test_size = len(X_test)
#     unique_labels = np.unique(y_test)

#     print(f"Số lượng phần tử trong tập test: {test_size}")
#     print(f"Giá trị nhãn khác nhau trong tập test: {unique_labels}")

#     doc.add_paragraph(f"Số lượng phần tử trong tập test: {test_size}")
#     doc.add_paragraph(f"Giá trị nhãn khác nhau trong tập test: {unique_labels}")

#     # Mô hình Naive Bayes
#     gnb = GaussianNB()
#     gnb.fit(X_train, y_train)
#     y_pred_bayes = gnb.predict(X_test)

#     # Đánh giá Naive Bayes
#     accuracy_bayes = accuracy_score(y_test, y_pred_bayes) * 100
#     report = classification_report(y_test, y_pred_bayes, zero_division=1)

#     print(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     print("Báo cáo phân loại Naive Bayes:")
#     print(report)

#     doc.add_paragraph(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%")
#     doc.add_paragraph("Báo cáo phân loại Naive Bayes:")
#     doc.add_paragraph(report)

#     # Lưu kết quả vào file TXT với mã hóa UTF-8
#     with open(txt_file_path, 'a', encoding='utf-8') as f:
#         f.write(f"Số lượng phần tử trong tập test: {test_size}\n")
#         f.write(f"Giá trị nhãn khác nhau trong tập test: {unique_labels}\n")
#         f.write(f"Độ chính xác tổng thể của Naive Bayes: {accuracy_bayes:.2f}%\n")
#         f.write("Báo cáo phân loại Naive Bayes:\n")
#         f.write(report)
#         f.write("\n" + "="*50 + "\n")  # Để phân cách giữa các lần phân chia

#     break  # Chỉ chạy lần phân chia đầu tiên

# # Lưu file DOCX
# doc.save(docx_file_path)
# print(f"Kết quả đã được lưu vào file: {docx_file_path}")
# print(f"Kết quả đã được lưu vào file: {txt_file_path}")
